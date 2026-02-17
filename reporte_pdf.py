import os
import io
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import pythoncom

def set_cell_shading(cell, fill_color):
    """Establece el color de fondo de una celda (Formato Hexadecimal sin #)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def generar_pdf_estado_cuenta(cliente_data, cartera_df):
    """
    Genera el estado de cuenta profesional de ALPAPEL SAS.
    Configura Encabezados y Pies de página para que los logos se repitan en cada hoja.
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    exports_dir = os.path.join(base_dir, "exports")
    if not os.path.exists(exports_dir):
        os.makedirs(exports_dir)

    timestamp = datetime.now().strftime("%H%M%S")
    temp_docx = os.path.join(exports_dir, f"estado_cuenta_{timestamp}.docx")
    temp_pdf = temp_docx.replace(".docx", ".pdf")

    # 1. Crear documento y configurar fuente base
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # --- CONFIGURACIÓN DE SECCIÓN (Encabezado y Pie de Página) ---
    section = doc.sections[0]
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # 2. ENCABEZADO (Logo Izquierda + Título Derecha Centrado)
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7))
    
    # Logo superior en el encabezado
    logo_path = os.path.join(base_dir, "assets", "Logo_formato.jpg")
    if os.path.exists(logo_path):
        logo_cell = header_table.rows[0].cells[0]
        p_logo = logo_cell.paragraphs[0]
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(2.0))
    
    # Título en el encabezado (Lado derecho, centrado internamente)
    title_cell = header_table.rows[0].cells[1]
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ESTADO DE CUENTA\nALPAPEL SAS")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 179, 176) # Turquesa ALPAPEL

    # 3. PIE DE PÁGINA (Logo Inferior)
    footer = section.footer
    logo_inf_path = os.path.join(base_dir, "assets", "Logo_inferior.jpg")
    if os.path.exists(logo_inf_path):
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = p_footer.add_run()
        run_footer.add_picture(logo_inf_path, width=Inches(6.0))

    # --- CUERPO DEL DOCUMENTO ---

    # 4. Datos del Cliente
    doc.add_paragraph() # Espacio inicial para no pegar al encabezado
    nombre = (cliente_data.get('nombre_cliente') or cliente_data.get('razon_social') or 'CLIENTE').upper()
    nit = str(cliente_data.get('nit_cliente') or 'N/A')
    fecha_emision = datetime.now().strftime("%d/%m/%Y %I:%M %p")

    def add_info_line(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    add_info_line("CLIENTE", nombre)
    add_info_line("NIT/C.C", nit)
    add_info_line("FECHA EMISIÓN", fecha_emision)

    # 5. Procesar Cartera
    lista_facturas_todas = []
    total_mora = 0
    total_cartera = 0

    for _, row in cartera_df.iterrows():
        monto = 0
        for col in ['saldo', 'total_cop']:
            if col in row and pd.notna(row[col]):
                monto = float(row[col])
                break
        
        total_cartera += monto
        dias_mora = int(row.get('dias_vencidos', 0))
        if dias_mora > 0:
            total_mora += monto

        lista_facturas_todas.append({
            'doc': str(row.get('documento') or row.get('nro_factura') or 'N/A'),
            'emision': str(row.get('fecha_emision', 'N/A')).split(' ')[0],
            'vence': str(row.get('fecha_vencimiento', 'N/A')).split(' ')[0],
            'mora': str(dias_mora) if dias_mora > 0 else "0",
            'monto': f"${monto:,.0f}"
        })

    doc.add_paragraph()

    # 6. Tabla corporativa (100% de facturas)
    doc.add_paragraph().add_run("DETALLE DE DOCUMENTOS").bold = True
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    titulos = ['DOCUMENTO', 'EMISIÓN', 'VENCIMIENTO', 'DÍAS MORA', 'SALDO']
    
    for i, titulo in enumerate(titulos):
        cell = table.rows[0].cells[i]
        cell.text = titulo
        set_cell_shading(cell, "00B3B0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True

    for f in lista_facturas_todas:
        row_cells = table.add_row().cells
        row_cells[0].text = f['doc']
        row_cells[1].text = f['emision']
        row_cells[2].text = f['vence']
        row_cells[3].text = f['mora']
        row_cells[4].text = f['monto']
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if int(f['mora']) > 0:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # 7. Totales
    doc.add_paragraph()
    p_totales = doc.add_paragraph()
    p_totales.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_mora = p_totales.add_run(f"TOTAL EN MORA: ${total_mora:,.0f}\n")
    run_mora.bold = True
    run_mora.font.color.rgb = RGBColor(200, 0, 0)
    
    run_total = p_totales.add_run(f"TOTAL SALDO CARTERA: ${total_cartera:,.0f}")
    run_total.bold = True
    run_total.font.size = Pt(11)

    # 8. Nota Final
    doc.add_paragraph()
    doc.add_paragraph("___________________________________________________________________________")
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_nota = p_nota.add_run(
        "Nota: Este extracto refleja la totalidad de facturas pendientes en nuestro sistema (al día y vencidas). "
        "Si usted ya realizó el pago, por favor envíe su soporte a quien le ha hecho llegar este documento o a los correos: "
        "tesoreria@alpapel.com / coordinador.cartera@alpapel.com / cartera@alpapel.com para su debida legalización."
    )
    run_nota.font.italic = True
    run_nota.font.size = Pt(8)
    run_nota.font.color.rgb = RGBColor(80, 80, 80)

    # 9. Guardar y Convertir
    doc.save(temp_docx)
    try:
        pythoncom.CoInitialize()
        convert(temp_docx, temp_pdf)
        with open(temp_pdf, "rb") as f:
            pdf_data = f.read()
        buffer = io.BytesIO(pdf_data)
        buffer.seek(0)
        return buffer
    finally:
        if os.path.exists(temp_docx): os.remove(temp_docx)
        if os.path.exists(temp_pdf): os.remove(temp_pdf)